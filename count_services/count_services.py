import docx
import pickle
import re
import os
from . import config

def remove_chars(line, chars):
    return re.sub(r'[' + chars + r']', "", line)


def remove_words(line, words):
    result = ''
    for word in line.split():
        if word not in words:
            result += word + ' '
    return result.rstrip()


class ServiceCounter:

    def __init__(self, k):
        self.k = k
        try:
            with open(config.KEYWORDS_FILE, 'rb') as f:
                self.keywords = pickle.load(f)
        except:
            self.keywords = set()

    def train(self):
        for filename, services_info in config.CONTRACT_DOCS.items():
            contract = docx.Document(filename)
            tables = contract.tables
            self.keywords = set()

            # run through the tables
            for table_idx, column_idx in zip(services_info[0], services_info[1]):
                table = tables[table_idx]
                for cell in table.column_cells(column_idx)[1:]:
                    operations = cell.text.lower().split('\n')

                    for oper in operations:
                        oper = oper.split()

                        for word in oper:
                            word = remove_chars(word, config.PUNCTUATION)
                            if word and word not in config.STOP_WORDS and not word.isdigit():
                                self.keywords.add(word)
        with open(config.KEYWORDS_FILE, 'wb') as f:
            pickle.dump(self.keywords, f)	


    def predict_proba(self, sentance):
        overlap = 0

        sentance = remove_chars(remove_words(sentance, config.STOP_WORDS), config.PUNCTUATION).lower().split()
        for word in sentance:
            if word in self.keywords:
                overlap += 1
        return overlap / len(sentance) if sentance else 0


    def count_services(self, contract):
        counter = 0
        current_path = os.path.dirname(__file__)
        keyword_path = os.path.join(current_path, config.KEYWORDS_FILE)
        with open(keyword_path, 'rb') as f:
            self.keywords = pickle.load(f)

        tables = contract.tables

        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    sentances = cell.text.lower().split('\n')
                    for sentance in sentances:
                        if self.predict_proba(sentance) >= self.k:
                            counter += 1
        return counter


if __name__ == "__main__":
    counter = ServiceCounter(0.7)
    counter.train()
    # counter.keywords = {"мойка"}
    # print(counter.keywords)
    print(counter.predict_proba("Тут вообще нет ни слова про услугу"))
    print(counter.predict_proba("Мойка полов в помещении"))
    # # with open(config.KEYWORDS_FILE, 'rb') as f:
    # # 	keywords = pickle.load(f)
    # # print(keywords)
